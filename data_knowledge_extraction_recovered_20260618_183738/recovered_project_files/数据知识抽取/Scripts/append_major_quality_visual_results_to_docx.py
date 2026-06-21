from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "数据知识挖掘.docx"
TABLE_VERSION_BACKUP = ROOT / "docs" / "word_backups" / "数据知识挖掘.before_visual_replacement_20260612.docx"
BASE_BEFORE_RESULTS = ROOT / "docs" / "word_backups" / "数据知识挖掘.before_major_quality_results_20260612.docx"
BACKUP_DIR = ROOT / "docs" / "word_backups"
RESULT_DIR = ROOT / "knowledge_exports" / "major_quality_abnormality_experiment_v1"
FIGURE_DIR = RESULT_DIR / "visual_docx_figures"
RESULT_JSON = RESULT_DIR / "major_quality_abnormality_experiment_seed42.json"
BINARY_IMPORTANCE = RESULT_DIR / "binary_feature_importance_seed42.csv"
MULTILABEL_IMPORTANCE = RESULT_DIR / "multilabel_feature_importance_seed42.csv"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)

PALETTE = {
    "blue": (54, 111, 196),
    "teal": (32, 145, 140),
    "green": (67, 160, 71),
    "orange": (245, 124, 0),
    "red": (211, 47, 47),
    "purple": (126, 87, 194),
    "gray": (110, 120, 130),
    "light_grid": (228, 232, 238),
    "text": (35, 43, 55),
    "muted": (100, 112, 128),
    "bg": (255, 255, 255),
}


def font(size: int, *, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def draw_text_fit(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font_obj: ImageFont.ImageFont,
    *,
    fill: tuple[int, int, int] = PALETTE["text"],
    max_width: int | None = None,
) -> None:
    if max_width is None or draw.textlength(text, font=font_obj) <= max_width:
        draw.text(xy, text, font=font_obj, fill=fill)
        return
    clipped = text
    while clipped and draw.textlength(clipped + "...", font=font_obj) > max_width:
        clipped = clipped[:-1]
    draw.text(xy, clipped + "...", font=font_obj, fill=fill)


def fmt(value: Any, digits: int = 3) -> str:
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def canvas(title: str, subtitle: str = "", *, width: int = 1500, height: int = 900) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (width, height), PALETTE["bg"])
    draw = ImageDraw.Draw(img)
    draw.text((60, 42), title, font=font(40, bold=True), fill=PALETTE["text"])
    if subtitle:
        draw.text((60, 96), subtitle, font=font(24), fill=PALETTE["muted"])
    draw.line((60, 138, width - 60, 138), fill=PALETTE["light_grid"], width=2)
    return img, draw


def save_horizontal_bars(
    path: Path,
    title: str,
    subtitle: str,
    labels: list[str],
    values: list[float],
    *,
    value_suffix: str = "",
    max_value: float | None = None,
    colors: list[tuple[int, int, int]] | None = None,
    width: int = 1500,
    height: int = 900,
) -> None:
    img, draw = canvas(title, subtitle, width=width, height=height)
    left_label = 70
    left_bar = 480
    right = width - 130
    top = 190
    row_gap = 72
    bar_h = 34
    max_v = max_value if max_value is not None else max(values) * 1.08
    colors = colors or [PALETTE["blue"]] * len(labels)
    axis_font = font(22)
    value_font = font(22, bold=True)

    for idx, (label, value) in enumerate(zip(labels, values)):
        y = top + idx * row_gap
        draw_text_fit(draw, (left_label, y - 3), label, axis_font, max_width=380)
        draw.rounded_rectangle((left_bar, y, right, y + bar_h), radius=10, fill=(239, 242, 247))
        bar_w = int((right - left_bar) * (float(value) / max_v)) if max_v else 0
        draw.rounded_rectangle((left_bar, y, left_bar + bar_w, y + bar_h), radius=10, fill=colors[idx % len(colors)])
        if value_suffix == "%":
            value_text = f"{value * 100:.1f}%"
        elif value_suffix:
            value_text = f"{value:,.0f}{value_suffix}"
        else:
            value_text = f"{value:.3f}" if value <= 1.2 else f"{value:,.0f}"
        draw.text((right + 18, y - 2), value_text, font=value_font, fill=PALETTE["text"])

    img.save(path, quality=95)


def save_grouped_class_metrics(path: Path, per_label: dict[str, Any]) -> None:
    labels = [item["display_name"] for item in per_label.values()]
    precision = [float(item["precision"]) for item in per_label.values()]
    recall = [float(item["recall"]) for item in per_label.values()]
    f1 = [float(item["f1"]) for item in per_label.values()]
    ap = [float(item["average_precision"]) for item in per_label.values()]

    img, draw = canvas("五类主要异常识别效果", "按类别展示 Precision / Recall / F1 / AP", width=1600, height=960)
    left = 110
    right = 1510
    bottom = 760
    top = 190
    chart_h = bottom - top
    chart_w = right - left
    max_v = 1.0
    grid_font = font(20)
    label_font = font(18)
    legend_font = font(22)
    series = [
        ("Precision", precision, PALETTE["blue"]),
        ("Recall", recall, PALETTE["teal"]),
        ("F1", f1, PALETTE["orange"]),
        ("AP", ap, PALETTE["purple"]),
    ]
    for tick in [0.0, 0.25, 0.50, 0.75, 1.0]:
        y = bottom - int(chart_h * tick / max_v)
        draw.line((left, y, right, y), fill=PALETTE["light_grid"], width=1)
        draw.text((55, y - 12), f"{tick:.2f}", font=grid_font, fill=PALETTE["muted"])
    group_w = chart_w / len(labels)
    bar_w = 42
    for group_idx, label in enumerate(labels):
        center = left + int(group_w * group_idx + group_w / 2)
        for s_idx, (_, values, color) in enumerate(series):
            x0 = center - 2 * bar_w + s_idx * bar_w
            value = values[group_idx]
            y0 = bottom - int(chart_h * value / max_v)
            draw.rounded_rectangle((x0, y0, x0 + bar_w - 6, bottom), radius=5, fill=color)
        draw_text_fit(draw, (center - 115, bottom + 24), label, label_font, max_width=230)
    legend_x = 350
    legend_y = 820
    for name, _, color in series:
        draw.rounded_rectangle((legend_x, legend_y, legend_x + 26, legend_y + 26), radius=6, fill=color)
        draw.text((legend_x + 38, legend_y - 2), name, font=legend_font, fill=PALETTE["text"])
        legend_x += 250
    img.save(path, quality=95)


def save_feature_panels(path: Path, binary_imp: pd.DataFrame, multilabel_imp: pd.DataFrame) -> None:
    img, draw = canvas("关键工艺特征重要性", "左：二分类风险识别；右：五类多标签平均重要性", width=1700, height=1020)
    panel_specs = [
        ("主要异常二分类", binary_imp.head(10), "binary_importance", 70, 820, PALETTE["blue"]),
        ("五类多标签平均", multilabel_imp.head(10), "multilabel_mean_importance", 900, 1650, PALETTE["teal"]),
    ]
    title_font = font(28, bold=True)
    label_font = font(18)
    value_font = font(17, bold=True)
    for panel_title, df, col, left, right, color in panel_specs:
        draw.text((left, 170), panel_title, font=title_font, fill=PALETTE["text"])
        max_v = float(df[col].max()) * 1.10
        top = 230
        row_gap = 66
        label_w = 285
        bar_left = left + label_w
        for idx, row in df.iterrows():
            y = top + len(df.loc[:idx]) * 0
        for rank, (_, row) in enumerate(df.iterrows()):
            y = top + rank * row_gap
            feature = str(row["feature"])
            value = float(row[col])
            draw_text_fit(draw, (left, y - 1), feature, label_font, max_width=label_w - 20)
            draw.rounded_rectangle((bar_left, y, right, y + 30), radius=8, fill=(239, 242, 247))
            bar_w = int((right - bar_left) * value / max_v) if max_v else 0
            draw.rounded_rectangle((bar_left, y, bar_left + bar_w, y + 30), radius=8, fill=color)
            draw.text((right + 12, y - 1), f"{value:.4f}", font=value_font, fill=PALETTE["text"])
    img.save(path, quality=95)


def save_figures(result: dict[str, Any]) -> dict[str, Path]:
    FIGURE_DIR.mkdir(parents=True, exist_ok=True)
    binary = result["binary"]
    multilabel = result["multilabel"]
    binary_imp = pd.read_csv(BINARY_IMPORTANCE)
    multilabel_imp = pd.read_csv(MULTILABEL_IMPORTANCE)

    paths = {
        "dataset": FIGURE_DIR / "01_dataset_composition.png",
        "binary": FIGURE_DIR / "02_binary_metrics.png",
        "overall": FIGURE_DIR / "03_multilabel_overall_metrics.png",
        "class": FIGURE_DIR / "04_per_class_metrics.png",
        "features": FIGURE_DIR / "05_feature_importance.png",
    }
    save_horizontal_bars(
        paths["dataset"],
        "主要异常数据构成",
        "主标签、干净负样本与上下文辅助样本的规模关系",
        ["主要异常样本", "干净负样本", "辅助状态-only样本", "主异常多标签样本"],
        [45185, 30187, 23374, 9742],
        colors=[PALETTE["blue"], PALETTE["green"], PALETTE["gray"], PALETTE["orange"]],
    )
    save_horizontal_bars(
        paths["binary"],
        "主要异常二分类结果",
        "主要异常 vs 干净负样本，group split by process_sequence_key",
        ["F1", "AUC", "AP", "Accuracy", "Precision", "Recall", "Balanced Accuracy"],
        [
            float(binary["f1"]),
            float(binary["roc_auc"]),
            float(binary["average_precision"]),
            float(binary["accuracy"]),
            float(binary["precision"]),
            float(binary["recall"]),
            float(binary["balanced_accuracy"]),
        ],
        value_suffix="%",
        max_value=1.0,
        colors=[PALETTE["blue"], PALETTE["purple"], PALETTE["teal"], PALETTE["gray"], PALETTE["green"], PALETTE["orange"], PALETTE["red"]],
    )
    save_horizontal_bars(
        paths["overall"],
        "五类主要异常多标签总体结果",
        "多标签识别的总体指标",
        ["Micro-F1", "Macro-F1", "mAP", "Samples-F1", "Subset Accuracy"],
        [
            float(multilabel["micro_f1"]),
            float(multilabel["macro_f1"]),
            float(multilabel["mean_average_precision"]),
            float(multilabel["sample_f1"]),
            float(multilabel["subset_accuracy"]),
        ],
        value_suffix="%",
        max_value=1.0,
        colors=[PALETTE["blue"], PALETTE["teal"], PALETTE["purple"], PALETTE["orange"], PALETTE["gray"]],
    )
    save_grouped_class_metrics(paths["class"], multilabel["per_label"])
    save_feature_panels(paths["features"], binary_imp, multilabel_imp)
    return paths


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def setup_styles(doc: Document) -> None:
    for style_name in ("Normal", "List Bullet"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run)


def add_bullet(doc: Document, text: str) -> None:
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
        prefix = ""
    except KeyError:
        paragraph = doc.add_paragraph()
        prefix = "- "
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(prefix + text)
    set_run(run, size=10.5)


def add_picture_with_caption(doc: Document, path: Path, caption: str, *, width: float = 6.55) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_p.add_run(caption)
    set_run(caption_run, size=9.5, color=RGBColor(88, 96, 105))


def backup_current_table_version() -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    if DOCX.exists() and not TABLE_VERSION_BACKUP.exists():
        shutil.copy2(DOCX, TABLE_VERSION_BACKUP)
    return TABLE_VERSION_BACKUP


def append_visual_section(doc: Document, result: dict[str, Any], figures: dict[str, Path]) -> None:
    today = date.today().isoformat()
    doc.add_page_break()
    add_heading(doc, f"主要异常数据处理与模型基线结果（图形化版，{today}）", 1)
    add_para(
        doc,
        "本节将主要异常处理结果和模型基线结果改为图形化呈现。当前口径把主标签限定为液面/卷渣质量类、温度/保护剂质量类、传热不均质量类、拉速-塞棒/流量控制异常类、过程参数异常类；交接过渡/质量状态类只作为上下文辅助变量保留，不进入主标签。",
    )
    add_bullet(doc, "模型采用 ExtraTreesClassifier，80 棵树，随机种子 42，中位数填补缺失值。")
    add_bullet(doc, "切分方式采用 process_sequence_key 分组切分，降低同一工况片段同时出现在训练集和测试集的风险。")
    add_bullet(doc, "特征策略为仅使用 99 个工艺数值特征，排除品质异常代码、标签字段、改钢原因、处置代码、类别字段和样本 ID 字段。")

    add_heading(doc, "1. 数据构成", 2)
    add_picture_with_caption(doc, figures["dataset"], "图1 主要异常数据构成：主异常样本、干净负样本、辅助状态样本与多标签样本规模。")

    add_heading(doc, "2. 二分类风险识别结果", 2)
    binary = result["binary"]
    add_para(
        doc,
        f"主要异常二分类达到 F1={fmt(binary['f1'])}、AUC={fmt(binary['roc_auc'])}、AP={fmt(binary['average_precision'])}。混淆矩阵为 TN=5,595、FP=310、FN=462、TP=8,708。",
    )
    add_picture_with_caption(doc, figures["binary"], "图2 主要异常二分类结果：风险识别在 F1、AUC、AP 和召回率上均较高。")

    add_heading(doc, "3. 五类主要异常多标签识别结果", 2)
    multilabel = result["multilabel"]
    add_para(
        doc,
        f"五类主要异常多标签识别达到 Micro-F1={fmt(multilabel['micro_f1'])}、Macro-F1={fmt(multilabel['macro_f1'])}、mAP={fmt(multilabel['mean_average_precision'])}。其中传热不均类测试支持数较少，后续应做多 seed 和置信区间验证。",
    )
    add_picture_with_caption(doc, figures["overall"], "图3 五类主要异常多标签总体结果。")
    add_picture_with_caption(doc, figures["class"], "图4 分异常类别的 Precision、Recall、F1 和 AP。")

    add_heading(doc, "4. 关键工艺特征", 2)
    add_para(
        doc,
        "特征重要性前列集中在过热度、TD 温度、液面波动、拉速波动、TD 吨位波动、南北拔热量差等工艺变量，与连铸品质异常机理基本一致。该结果说明当前模型主要利用工艺过程信号，而不是直接依赖品质异常编码或改钢原因字段。",
    )
    add_picture_with_caption(doc, figures["features"], "图5 关键工艺特征重要性：左侧为二分类风险识别，右侧为五类多标签平均重要性。")

    add_heading(doc, "5. 结果解释与下一步验证", 2)
    add_bullet(doc, "重新定义主要异常标签后，模型效果显著提升，说明此前效果差主要来自标签混杂和负样本不稳定。")
    add_bullet(doc, "当前结果仍需通过时间外推、连铸线外推、去除重复派生特征和多随机种子验证，确认其泛化能力。")
    add_bullet(doc, "后续应继续引入 rule_margin、非线性滞后特征和工艺先验路径遮蔽验证，用于证明规则边界、时序传播和路径解释的增益。")


def main() -> None:
    if not BASE_BEFORE_RESULTS.exists():
        raise FileNotFoundError(BASE_BEFORE_RESULTS)
    if not RESULT_JSON.exists():
        raise FileNotFoundError(RESULT_JSON)

    backup_current_table_version()
    result = json.loads(RESULT_JSON.read_text(encoding="utf-8"))
    figures = save_figures(result)

    shutil.copy2(BASE_BEFORE_RESULTS, DOCX)
    doc = Document(DOCX)
    setup_styles(doc)
    append_visual_section(doc, result, figures)
    doc.save(DOCX)

    print(f"updated={DOCX}")
    print(f"table_version_backup={TABLE_VERSION_BACKUP}")
    for name, path in figures.items():
        print(f"figure_{name}={path}")


if __name__ == "__main__":
    main()
