from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "kiepgl_data_model_optimization_report.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            cell = row.cells[idx]
            cell.width = Inches(width)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", style: str | None = None, bold_label: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_label:
        r = p.add_run(bold_label)
        set_run_font(r, bold=True)
        r = p.add_run(text)
        set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0]
    set_run_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK_BLUE)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    r = p.add_run("\n" + body)
    set_run_font(r)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_shading(cell, LIGHT_GRAY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx != 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_run_font(r, size=10.5)
    set_table_width(table, widths)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def build_doc() -> None:
    doc = Document()
    setup_styles(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = title.add_run("KIEP-GL 数据与模型优化报告")
    set_run_font(r, size=22, bold=True, color=DARK_BLUE)
    subtitle = doc.add_paragraph()
    r = subtitle.add_run("面向连铸规则边界与工艺耦合的边界区失稳风险预警")
    set_run_font(r, size=12, color=RGBColor(80, 80, 80))

    meta = doc.add_table(rows=4, cols=2)
    meta.style = "Table Grid"
    meta_rows = [
        ("项目阶段", "数据标签重构、负样本稳定化、事件组切分与模型适配"),
        ("推荐主任务", "正常窗口 vs 边界失稳风险窗口（二分类风险预警）"),
        ("推荐主配置", "risk_binary + horizon/stable normal + event-bag group split + flat KIEP-GL"),
        ("报告依据", "截至 2026-06-10 的本地实验结果与代码验证"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = v
        set_cell_shading(meta.cell(i, 0), LIGHT_GRAY)
    set_table_width(meta, [1.65, 4.85])
    doc.add_paragraph()

    add_callout(
        doc,
        "核心结论",
        "原始细粒度缺陷标签不适合作为主分类任务；将主任务调整为边界区失稳风险预警后，在事件组切分下 3000 条样本实验仍达到 target F1=0.8219，满足 0.7 以上目标。",
    )

    add_heading(doc, "1. 当前主要问题", 1)
    add_paragraph(
        doc,
        "前期实验表明，直接预测 temperature_flux、mold_level_slag_risk、process_fluctuation、speed_stopper_flow 等细粒度 subtype 时，模型的 target F1 长期停留在 0.23-0.29 区间。问题根源不是单纯模型容量不足，而是监督信号本身与工艺机理之间存在错位。",
    )
    add_bullet(doc, "标签边界不清：多缺陷窗口和单缺陷窗口混杂，细类之间并非互斥。")
    add_bullet(doc, "负样本定义不稳定：严格 normal 数量极少，过度放宽会引入短期正常但长期风险样本。")
    add_bullet(doc, "细类可分性不足：粗机制类实验只提升到 target F1=0.3173，说明现有过程特征更适合判断风险状态，而不是稳定区分所有缺陷类型。")
    add_bullet(doc, "窗口级切分有评估风险：同一事件的 early/mid/late 窗口可能跨训练和测试，因此必须改为 event-bag 级 group split。")

    add_heading(doc, "2. 数据优化方案", 1)
    add_table(
        doc,
        ["模块", "处理方式", "目的"],
        [
            ["稳定 normal", "启用 stable_normal_only，剔除物理无效、极端工况、脏质量角色样本", "减少负样本噪声"],
            ["阶段 horizon", "启用 stage_specific_negative_exclusion，late/mid/early 分别定义未来干净窗口", "避免一个 normal 定义混用所有预警阶段"],
            ["标签重构", "新增 risk_binary 标签：no_quality_abnormal vs boundary_instability_risk", "让主任务贴合边界失稳传播风险建模"],
            ["原始细类保留", "保存 original_event_quality_class_id 和 original_event_quality_class", "支持后续候选机制解释和路径归因"],
            ["事件组切分", "新增 group_stratified_time_ordered，按 event_bag_id 分组", "避免同一事件窗口跨 split 造成乐观评估"],
        ],
        [1.35, 3.15, 2.00],
    )

    add_heading(doc, "3. 标签方案对比", 1)
    add_table(
        doc,
        ["标签方案", "任务含义", "target F1", "event recall", "误报率", "判断"],
        [
            ["原细粒度 subtype", "预测具体缺陷类型", "0.2907", "0.2708", "0.1875", "不适合作主任务"],
            ["粗机制类", "热边界/流动控制/过程波动", "0.3173", "0.3111", "0.3000", "仍不足以稳定分类"],
            ["risk_binary", "边界区失稳风险预警", "0.8327", "0.8500", "0.1917", "推荐主任务"],
            ["risk_binary 低误报", "加入 normal-bias 校准", "0.7707", "0.6583", "0.0500", "适合保守部署"],
        ],
        [1.45, 1.85, 0.85, 0.9, 0.75, 1.00],
    )

    add_heading(doc, "4. 最新模型适配结果", 1)
    add_paragraph(
        doc,
        "在更大样本实验中，采用 3000 条 stratified 样本、event-bag group split、ExtraTrees 风险头、120 棵树和最大深度 12。该设置比 1200 条 smoke 更稳，且仍然保持 0.8 以上 F1。",
    )
    add_table(
        doc,
        ["实验", "样本", "切分", "target F1", "risk recall", "误报率", "平均提前量"],
        [
            ["主模型", "3000", "event-bag group split", "0.8219", "0.8000", "0.1519", "5.53 步"],
            ["主模型 smoke", "1200", "event-bag group split", "0.8327", "0.8500", "0.1917", "5.29 步"],
            ["低误报模式", "1200", "event-bag group split", "0.7707", "0.6583", "0.0500", "6.35 步"],
        ],
        [1.2, 0.7, 1.6, 0.8, 0.85, 0.75, 1.0],
    )

    add_heading(doc, "5. 推荐技术路线", 1)
    add_paragraph(
        doc,
        "建议将论文方法统一表述为“边界区失稳传播风险建模”，而不是“多缺陷细分类”。Temporal、Graph、rule_margin、KIEP-GL 路径解释都服务于同一主线：识别工艺变量接近规则边界后是否会在未来 horizon 内演化为质量风险。",
    )
    add_bullet(doc, "主任务：边界区失稳风险预警，报告 risk_binary F1、event recall、false alarm、lead time。")
    add_bullet(doc, "辅助任务：粗机制类和原始 subtype 仅作为候选解释，不作为主性能指标。")
    add_bullet(doc, "解释验证：报告 top-k path hit、path occlusion drop、event lead time、path consistency。")
    add_bullet(doc, "部署模式：主模型用于高召回预警，normal-bias 校准模型用于低误报场景。")

    add_heading(doc, "6. 代码与输出文件", 1)
    add_table(
        doc,
        ["对象", "路径/文件", "说明"],
        [
            ["标签重构脚本", "Scripts/build_kiepgl_label_scheme_dataset_v7.py", "生成 risk_binary 和 coarse_mechanism 数据集"],
            ["实验入口", "Scripts/run_kiepgl_multiclass_experiment.py", "支持 --split-mode group_stratified_time_ordered"],
            ["主数据集", "knowledge_exports/quality_traceability_dataset_v7_risk_binary_line_2_1_horizon_stable", "推荐用于论文主实验"],
            ["主实验结果", "knowledge_exports/kiepgl_v7_risk_binary_line_horizon_stable_group_split_flat_3000_t120d12.json", "3000 样本 group split 结果"],
            ["低误报结果", "knowledge_exports/kiepgl_v7_risk_binary_line_horizon_stable_group_split_flat_calibrated_smoke.json", "normal-bias 校准版本"],
        ],
        [1.35, 3.65, 1.5],
    )

    add_heading(doc, "7. 下一步建议", 1)
    add_bullet(doc, "扩展到 full-size event-bag group split，并补充多随机种子结果，报告均值和标准差。")
    add_bullet(doc, "增加事件级置信区间和按 late/mid/early 分阶段的 recall/lead time。")
    add_bullet(doc, "将细粒度 subtype 改为 candidate mechanism ranking，报告 top-k hit，而不是硬 top-1 accuracy。")
    add_bullet(doc, "在 Word/论文正文中明确 rule_margin、stable normal、horizon negative 的定义，避免审稿人质疑边界区划分任意。")

    section = doc.add_section(WD_SECTION.CONTINUOUS)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("KIEP-GL 数据与模型优化报告")
    set_run_font(r, size=9, color=RGBColor(100, 100, 100))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
    print(OUT)
