from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_CANDIDATES = [
    ROOT / "数据知识挖掘_整合更新.docx",
    ROOT / "数据知识挖掘.docx",
]
OUT = ROOT / "数据知识挖掘_任务与处理方案更新.docx"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def source_docx() -> Path:
    for path in SOURCE_CANDIDATES:
        if path.exists():
            return path
    raise FileNotFoundError("未找到可作为底稿的 Word 文档。")


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def setup_styles(doc: Document) -> None:
    for style_name in ("Normal", "List Bullet"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str, bold_label: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_label:
        r = p.add_run(bold_label)
        set_run(r, bold=True)
        r = p.add_run(text)
        set_run(r)
    else:
        r = p.add_run(text)
        set_run(r)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK)
    r = p.add_run("\n" + body)
    set_run(r)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, GRAY_FILL)
        margins(cell)
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run(r, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            margins(cell)
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run(r, size=10)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph()


def append_update(doc: Document) -> None:
    today = date.today().isoformat()
    doc.add_page_break()
    add_heading(doc, f"当前任务与处理方案更新（{today}）", 1)
    add_callout(
        doc,
        "当前阶段结论",
        "当前项目不宜继续把具体缺陷 subtype 作为唯一主任务。更科学的主线应定义为：面向规则边界与工艺耦合的边界区失稳传播风险建模；二分类负责判断是否进入品质异常风险状态，多标签负责描述异常现象集合，路径-原因对齐负责证明模型解释是否与改钢原因和工艺机理一致。",
    )

    add_heading(doc, "1. 当前任务定义", 2)
    add_table(
        doc,
        ["任务层级", "标签/输出", "核心含义", "建议评价指标"],
        [
            [
                "第一层：风险预警",
                "quality_core_event 或 quality_abnormal_event",
                "判断窗口是否进入品质异常风险状态；推荐以核心品质异常作为主二分类任务，避免过宽异常代码稀释边界失稳问题。",
                "Macro-F1、Risk-F1、Normal Recall、False Alarm、Event Recall、Lead Time",
            ],
            [
                "第二层：异常现象识别",
                "defect_set 与 defect_multilabel_*",
                "同一窗口允许存在多种异常现象，不强制只选一个主异常；这比只看品质异常代码1更符合现场记录方式。",
                "Samples-F1、Macro-F1、Top-k Hit、覆盖率",
            ],
            [
                "第三层：根因/路径解释",
                "模型路径机制 vs reason_mechanism_set",
                "将改钢原因处理为机制、阶段和变量组集合，评价模型输出路径是否能命中或软匹配人工原因。",
                "Path-Reason Hit@k、Soft Alignment、Path Occlusion Drop、Lead Time",
            ],
        ],
        [1.15, 1.35, 2.65, 1.45],
    )

    add_heading(doc, "2. 数据重构与标签处理方案", 2)
    add_bullet(doc, "宽口径 quality_abnormal_event：任意有效品质缺陷异常代码1-5非零即为风险样本；该口径适合做全量异常感知，但标签边界较宽，容易引入噪声。")
    add_bullet(doc, "核心口径 quality_core_event：仅保留与边界失稳传播更直接相关的核心异常组，例如温度/保护渣、液面卷渣、工艺波动、拉速/塞棒/流量、传热异常。")
    add_bullet(doc, "transition_tundish 与 other_quality_abnormal 不再作为主二分类正样本的核心来源；若只出现这两类，则进入辅助样本，用于后续扩展和泛化验证。")
    add_bullet(doc, "若品质异常标签为0但改钢原因存在明确风险证据，则标记为 conflict_quality_reason_negative，并从主训练集剔除，避免负样本定义不稳定。")
    add_bullet(doc, "单线分析优先选择连铸线2/1，采用 strict stable negative 构造高置信负样本；该方案能显著降低误报，但需要进一步提升召回。")

    add_heading(doc, "3. 当前关键数据集", 2)
    add_table(
        doc,
        ["数据集", "规模/分布", "用途", "备注"],
        [
            [
                "quality_abnormal_event_cleanconflict",
                "99,001行；正样本68,814，负样本30,187；冲突辅助样本1,564",
                "宽口径品质异常风险建模",
                "覆盖面广，但异常定义较宽",
            ],
            [
                "quality_core_event",
                "75,417行；正样本45,230，负样本30,187；辅助样本25,148",
                "核心品质异常主任务",
                "更贴合规则边界与工艺耦合主线",
            ],
            [
                "line_2_1_event_bag_strict_neg3",
                "26,397个窗口；风险25,585，正常812",
                "单线高置信负样本风险预警",
                "目前二分类最有效方案",
            ],
            [
                "line_2_1_risk_binary_strict_neg3",
                "26,397个样本；用于二分类训练",
                "当前推荐的主二分类数据集",
                "适合做保守预警基线",
            ],
        ],
        [1.85, 1.65, 1.45, 1.55],
    )

    add_heading(doc, "4. 当前二分类实验结果", 2)
    add_table(
        doc,
        ["实验版本", "Macro-F1", "Risk-F1", "Normal Recall", "False Alarm", "Event Recall", "结论"],
        [
            [
                "宽口径全线 mixed_neg6",
                "0.4730",
                "0.6609",
                "0.1833",
                "0.8167",
                "0.8967",
                "召回高但误报过高，说明负样本边界不稳",
            ],
            [
                "核心口径全线 mixed_neg6",
                "0.4723",
                "0.6517",
                "0.1933",
                "0.8067",
                "0.8967",
                "仅收窄正样本不足以解决误报",
            ],
            [
                "连铸线2/1 relaxed_neg3",
                "0.3752",
                "0.6636",
                "0.0467",
                "0.9533",
                "0.9100",
                "正常类几乎无法识别，不建议作为主方案",
            ],
            [
                "连铸线2/1 strict_neg3",
                "0.6941",
                "0.7557",
                "1.0000",
                "0.0000",
                "0.6377",
                "当前最佳；高精度、零误报，但召回偏保守",
            ],
        ],
        [1.35, 0.72, 0.72, 0.80, 0.78, 0.78, 1.85],
    )

    add_heading(doc, "5. 模型与代码处理方案", 2)
    add_table(
        doc,
        ["模块/文件", "当前作用", "后续处理方向"],
        [
            [
                "Scripts/build_quality_traceability_dataset.py",
                "构建行级质量追溯数据、宽/核心二分类标签、多标签异常集合和原因集合。",
                "继续完善 rule_margin、工况分层、冲突样本审计字段。",
            ],
            [
                "Scripts/build_kiepgl_event_bag_dataset_v5.py",
                "构建事件窗口，保留 defect_set、reason_mechanism_set、reason_confidence 等解释字段。",
                "增加多尺度窗口和边界前后对照窗口。",
            ],
            [
                "Scripts/build_kiepgl_multiclass_dataset_v4.py",
                "支持8类异常现象，包括 transition_tundish 与 other_quality_abnormal。",
                "第二层建议改为多标签输出，而不是互斥单分类。",
            ],
            [
                "Scripts/multiclass_event_evaluation.py",
                "新增路径与改钢原因的 loose/soft alignment 指标。",
                "形成 Path-Reason Hit@k、Soft Alignment、Occlusion Drop 的正式评价协议。",
            ],
        ],
        [1.85, 2.25, 2.40],
    )

    add_heading(doc, "6. 下一步优化路线", 2)
    add_para(
        doc,
        "技术路线可以概括为：先用工艺规则和质量异常记录重构更可靠的边界区风险标签，再通过单线高置信负样本建立保守但可信的风险预警基线，随后引入多标签异常现象识别和路径-原因对齐评价，把模型从“是否报警”扩展到“为什么报警、提前多久报警、解释是否符合改钢原因”。"
    )
    add_bullet(doc, "优先做多随机种子和时间切分复验，确认 strict_neg3 的稳定性，而不是只依赖一次 smoke 结果。")
    add_bullet(doc, "在保持低误报的基础上做召回增强：可采用两阶段模型，第一阶段高置信报警，第二阶段对近边界样本做召回扩展。")
    add_bullet(doc, "对第二层任务采用多标签学习，评价多个异常现象集合的覆盖情况，避免强行指定唯一主异常。")
    add_bullet(doc, "对第三层任务建立改钢原因词表到机制序号、阶段序号、变量组序号的映射，并用 Path-Reason Hit@k 与软匹配分数证明图路径可解释性。")


def main() -> None:
    src = source_docx()
    doc = Document(src)
    setup_styles(doc)
    append_update(doc)
    doc.save(OUT)

    check = Document(OUT)
    text = "\n".join(p.text for p in check.paragraphs)
    if "当前任务与处理方案更新" not in text:
        raise RuntimeError("结构检查失败：未找到新增章节标题。")
    if "连铸线2/1 strict_neg3" not in text:
        raise RuntimeError("结构检查失败：未找到关键实验结果。")
    print(f"source={src}")
    print(f"output={OUT}")
    print(f"paragraphs={len(check.paragraphs)}")
    print(f"tables={len(check.tables)}")


if __name__ == "__main__":
    main()
