from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]


BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"


def main_docx_path() -> Path:
    candidates = [
        p
        for p in ROOT.glob("*.docx")
        if not p.name.startswith("~$") and p.name != "kiepgl_data_model_optimization_report.docx"
    ]
    if not candidates:
        raise FileNotFoundError("No root-level docx file found.")
    preferred = [p for p in candidates if "数据知识" in p.name]
    return preferred[0] if preferred else candidates[0]


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run(r)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK)
    r = p.add_run("\n" + body)
    set_run(r)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, LIGHT)
        margins(cell)
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run(r, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            margins(cell)
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run(r, size=10)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph()


def has_existing_update(doc: Document) -> bool:
    marker = "八、最新数据质量与模型适配优化"
    return any(marker in p.text for p in doc.paragraphs)


def integrate() -> Path:
    path = main_docx_path()
    doc = Document(path)
    if has_existing_update(doc):
        print(f"Skipped existing update in {path}")
        return path

    for style_name in ("Normal", "List Bullet"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(11)

    add_heading(doc, "八、最新数据质量与模型适配优化（2026-06-10）", 1)
    add_callout(
        doc,
        "最新结论",
        "细粒度缺陷 subtype 不适合作为当前主任务；将主任务改为“正常窗口 vs 边界失稳风险窗口”的 risk_binary 后，在 event-bag group split 下 3000 条样本实验 target F1 达到 0.8219，满足 0.7 以上目标。",
    )
    add_para(
        doc,
        "本轮优化重点不是继续堆叠复杂模型，而是重构监督信号与评估协议，使任务定义更贴合“规则边界与工艺耦合导致的边界区失稳传播风险”。",
    )
    add_bullet(doc, "负样本定义：保留 horizon/stable normal，即按 late/mid/early 分阶段判断未来是否干净，并剔除物理无效、极端工况和脏质量角色样本。")
    add_bullet(doc, "标签定义：新增 risk_binary，将所有可追踪异常统一为 boundary_instability_risk，同时保留 original_event_quality_class_id 与 original_event_quality_class 作为解释字段。")
    add_bullet(doc, "评估协议：新增 group_stratified_time_ordered，按 event_bag_id 分组，避免同一质量事件的 early/mid/late 窗口跨训练集和测试集。")
    add_bullet(doc, "模型适配：主模型采用 flat KIEP-GL 风险预警；细粒度 subtype 不再作为主分类指标，改为候选机制解释和路径归因。")

    add_heading(doc, "九、标签方案与任务适配对比", 1)
    add_table(
        doc,
        ["标签方案", "任务含义", "target F1", "event recall", "误报率", "结论"],
        [
            ["原细粒度 subtype", "预测具体缺陷类型", "0.2907", "0.2708", "0.1875", "不适合作主任务"],
            ["粗机制类", "热边界/流动控制/过程波动", "0.3173", "0.3111", "0.3000", "仍不足以稳定分类"],
            ["risk_binary", "边界区失稳风险预警", "0.8327", "0.8500", "0.1917", "推荐主任务"],
            ["risk_binary 低误报", "加入 normal-bias 校准", "0.7707", "0.6583", "0.0500", "适合保守部署"],
        ],
        [1.25, 1.65, 0.75, 0.82, 0.72, 1.31],
    )
    add_para(
        doc,
        "上述结果说明，当前数据对“是否进入边界失稳风险状态”具有较强监督有效性，但对具体缺陷 subtype 的互斥分类监督不足。论文主任务应调整为边界区失稳风险预警，细类仅作为解释或候选机制输出。",
    )

    add_heading(doc, "十、最新正式实验结果", 1)
    add_table(
        doc,
        ["实验配置", "样本量", "切分方式", "target F1", "risk recall", "误报率", "平均提前量"],
        [
            ["主模型 t120d12", "3000", "event-bag group split", "0.8219", "0.8000", "0.1519", "5.53 步"],
            ["主模型 smoke", "1200", "event-bag group split", "0.8327", "0.8500", "0.1917", "5.29 步"],
            ["低误报模式", "1200", "event-bag group split", "0.7707", "0.6583", "0.0500", "6.35 步"],
        ],
        [1.35, 0.65, 1.45, 0.75, 0.78, 0.70, 0.82],
    )
    add_para(
        doc,
        "推荐将 3000 样本 group split 结果作为当前阶段主结果：target_defect_macro_f1=0.8219，event_macro_recall=0.8000，false_alarm_rate=0.1519，wrong_class_alarm_rate=0.0。该结果比 1200 条 smoke 更稳，更适合写入论文实验章节。",
    )

    add_heading(doc, "十一、对应代码和数据文件", 1)
    add_table(
        doc,
        ["对象", "路径", "作用"],
        [
            ["标签重构脚本", "Scripts/build_kiepgl_label_scheme_dataset_v7.py", "生成 risk_binary 与 coarse_mechanism 数据集"],
            ["实验入口", "Scripts/run_kiepgl_multiclass_experiment.py", "支持 --split-mode group_stratified_time_ordered"],
            ["推荐数据集", "knowledge_exports/quality_traceability_dataset_v7_risk_binary_line_2_1_horizon_stable", "论文主实验数据"],
            ["主实验结果", "knowledge_exports/kiepgl_v7_risk_binary_line_horizon_stable_group_split_flat_3000_t120d12.json", "3000 样本正式结果"],
            ["低误报结果", "knowledge_exports/kiepgl_v7_risk_binary_line_horizon_stable_group_split_flat_calibrated_smoke.json", "保守校准模式"],
        ],
        [1.20, 3.70, 1.60],
    )

    add_heading(doc, "十二、论文表述建议", 1)
    add_para(doc, "建议将方法主线统一为：面向连铸规则边界与工艺耦合的边界区失稳传播风险建模。")
    add_bullet(doc, "主性能指标：risk_binary 的 target F1、event recall、false alarm rate 和 lead time。")
    add_bullet(doc, "解释性指标：KIEP-GL 路径命中率、路径遮挡下降、事件提前量和路径一致性。")
    add_bullet(doc, "细粒度缺陷：不作为主分类指标，而作为候选机制集合、top-k path hit 和案例解释。")
    add_bullet(doc, "部署口径：高召回主模型用于预警，normal-bias 校准模型用于低误报场景。")

    add_heading(doc, "十三、更新后的当前结论", 1)
    add_para(
        doc,
        "当前最可靠的技术路线是 risk_binary 边界失稳风险预警，而不是细粒度缺陷硬分类。通过稳定 normal、分阶段 horizon 负样本、event-bag group split 和风险标签重构，模型在更严格评估下仍能达到 0.8219 的 target F1。后续优化应集中在 full-size 多随机种子实验、事件级置信区间、分阶段 lead-time 统计和候选机制解释，而不是继续强行提升 subtype top-1 分类。",
    )

    doc.save(path)
    return path


if __name__ == "__main__":
    print(integrate().resolve())
