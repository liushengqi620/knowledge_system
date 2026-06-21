from __future__ import annotations

import shutil
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "数据知识挖掘_整合更新.docx"
FALLBACK_SOURCE = ROOT / "数据知识挖掘.docx"
OUT = ROOT / "数据知识挖掘.docx"
BACKUP_DIR = ROOT / "docs" / "word_backups"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def choose_source() -> Path:
    if SOURCE.exists():
        return SOURCE
    if FALLBACK_SOURCE.exists():
        return FALLBACK_SOURCE
    raise FileNotFoundError("未找到可用的 Word 底稿。")


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def setup_styles(doc: Document) -> None:
    for style_name in ("Normal", "List Bullet"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run(r, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r)


def add_bullet(doc: Document, text: str) -> None:
    try:
        p = doc.add_paragraph(style="List Bullet")
        prefix = ""
    except KeyError:
        p = doc.add_paragraph()
        prefix = "• "
    r = p.add_run(prefix + text)
    set_run(r)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=DARK)
    r = p.add_run("\n" + body)
    set_run(r)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, GRAY_FILL)
        margins(cell)
        cell.width = Inches(widths[i])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run(r, size=10.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            margins(cell)
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i in (0, 2, 3) else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(value)
            set_run(r, size=9.5)
    for row in table.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = Inches(width)
    doc.add_paragraph()


def backup_existing_output() -> None:
    if not OUT.exists():
        return
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"数据知识挖掘.before_three_dataset_update_{date.today().strftime('%Y%m%d')}.docx"
    if not backup.exists():
        shutil.copy2(OUT, backup)


def append_three_dataset_plan(doc: Document) -> None:
    today = date.today().isoformat()
    doc.add_page_break()
    add_heading(doc, f"当前任务重构与三数据集方案（{today}）", 1)
    add_callout(
        doc,
        "核心调整",
        "三层任务不再共用同一批样本，而是拆成三个数据集：风险预警数据集、异常现象多标签数据集、根因追溯/路径解释数据集。这样可以避免二分类、异常识别和根因解释三种监督信号互相污染，也能更清楚地回应审稿人关于标签边界、负样本稳定性和解释可验证性的质疑。",
    )

    add_heading(doc, "1. 为什么需要拆成三个数据集", 2)
    add_para(
        doc,
        "当前数据的主要矛盾不是模型容量不足，而是不同任务依赖的监督依据不同。第一层风险预警依赖品质异常是否发生；第二层异常现象识别依赖多个品质异常代码共同出现的集合；第三层根因追溯依赖改钢原因是否能提供可解释的工艺原因。若三层任务强行使用同一数据集，标签会同时受到品质代码、改钢原因、现场记录完整性和工况切换的影响，导致模型目标不清晰、负样本定义不稳定、路径解释难以定量验证。",
    )

    add_heading(doc, "2. 三层任务对应的数据集定义", 2)
    add_table(
        doc,
        ["任务", "推荐数据集", "样本进入规则", "排除/辅助规则", "主要用途"],
        [
            [
                "第一层：风险预警",
                "Risk-Binary Dataset",
                "以核心品质异常为正样本；负样本必须是未来窗口内无品质异常、无明确改钢原因证据的 strict stable normal。当前优先使用连铸线2/1 strict_neg3 版本作为高置信基线。",
                "剔除品质异常为0但改钢原因提示风险的冲突负样本；transition_tundish-only 和 other-only 样本进入辅助集，不作为主正样本。",
                "回答是否进入边界区失稳传播风险状态，评价 Macro-F1、Risk-F1、误报率、提前量。",
            ],
            [
                "第二层：异常现象识别",
                "Defect Multi-label Dataset",
                "使用所有有效品质异常代码1-5构造 defect_set，不把品质异常代码1视为唯一主异常；同一窗口可同时标注 temperature_flux、mold_level_slag_risk、process_fluctuation 等多个标签。",
                "物理无效样本剔除；记录模糊但存在异常的样本可保留为弱标签或低置信样本；正常样本只用于阈值校准和误报评估。",
                "回答发生了哪些异常现象，评价 Samples-F1、Macro-F1、Top-k 命中和异常集合覆盖率。",
            ],
            [
                "第三层：根因追溯/路径解释",
                "Root-cause Alignment Dataset",
                "只使用两类核心样本：未改钢稳定样本作为背景/无根因对照；3xx、4xx 改钢原因样本作为可解释根因样本，并映射到机制、阶段、变量组三类原因标签。",
                "排除无改钢原因、原因文本过于模糊、非工艺原因、1xx/2xx/5xx 等暂不纳入主根因评价的样本；这些样本可作为外部泛化或案例分析。",
                "回答模型路径是否能对齐改钢原因，评价 Path-Reason Hit@k、Soft Alignment、Path Occlusion Drop 和路径一致性。",
            ],
        ],
        [1.10, 1.20, 2.20, 2.00, 1.70],
    )

    add_heading(doc, "3. 根因追溯数据集的科学限定", 2)
    add_bullet(doc, "正样本不直接由品质异常代码决定，而由 3xx/4xx 改钢原因提供根因监督信号；这样第三层评价的是解释路径是否符合工艺处置原因，而不是重复预测品质异常。")
    add_bullet(doc, "未改钢样本不是普通负样本，而是背景/无处置对照样本；它用于检验模型是否会在没有处置依据时产生过度解释。")
    add_bullet(doc, "3xx/4xx 样本需要进一步映射为 reason_mechanism_set、reason_stage_set 和 reason_variable_group_set，例如温度-热流、液面-卷渣、拉速/塞棒-流动控制、结晶器/二冷-传热异常等。")
    add_bullet(doc, "根因任务的评价应采用宽松但可量化的匹配：只要模型 Top-k 路径命中同一机制、同一工艺阶段或同一变量组，即可计入 loose hit；完全一致则计入 strict hit。")

    add_heading(doc, "4. 三个数据集之间的关系", 2)
    add_para(
        doc,
        "三个数据集应共享同一套基础清洗、变量字典、时间窗口和 rule_margin 定义，但不共享同一套样本筛选逻辑。风险预警数据集追求正负边界稳定；异常现象数据集追求异常集合覆盖；根因追溯数据集追求原因标签可信。最终模型可以采用级联方式：第一层先判断是否预警，第二层输出可能的异常现象集合，第三层只在存在可解释原因证据或需要决策辅助时输出路径解释。",
    )
    add_table(
        doc,
        ["共享部分", "不共享部分", "这样处理的好处"],
        [
            [
                "基础清洗、物理有效性过滤、变量标准化、工况字段、rule_margin、时间窗口生成逻辑",
                "正负样本定义、标签来源、辅助样本规则、评价指标",
                "既保证三个任务来自同一工艺系统，又避免标签语义混杂。",
            ],
            [
                "事件窗口 event_bag_id、连铸线/炉次/时间信息、变量组映射",
                "二分类风险标签、多标签异常集合、3xx/4xx 根因标签",
                "可以在论文中形成统一框架：边界区失稳传播风险建模与可解释追溯。",
            ],
        ],
        [2.30, 2.25, 2.45],
    )

    add_heading(doc, "5. 对当前代码和实验的调整建议", 2)
    add_table(
        doc,
        ["代码/数据对象", "当前状态", "建议调整"],
        [
            [
                "Scripts/build_quality_traceability_dataset.py",
                "已经支持 quality_abnormal_event、quality_core_event、defect_set、reason_mechanism_set 等字段。",
                "继续保留为统一基础数据构建脚本，但输出三个任务专用 manifest，明确每个数据集的进入/排除规则。",
            ],
            [
                "Scripts/build_kiepgl_event_bag_dataset_v5.py",
                "已经能把 defect_set 和 reason_* 字段传递到窗口级数据。",
                "增加 task_mode：risk_binary、defect_multilabel、root_cause_alignment，避免下游实验混用样本。",
            ],
            [
                "Scripts/multiclass_event_evaluation.py",
                "已经加入路径与原因集合的 loose/soft alignment 指标。",
                "把 3xx/4xx 根因数据集作为路径解释评价主集，未改钢样本用于过度解释惩罚。",
            ],
            [
                "当前 Word 文档",
                "之前把三层任务写成一个整体数据方案。",
                "改为三数据集方案，并把 risk、defect、root-cause 三个任务的边界分别说明。",
            ],
        ],
        [1.75, 2.35, 2.90],
    )

    add_heading(doc, "6. 推荐写入论文的方法主线", 2)
    add_para(
        doc,
        "论文主线建议表述为：针对连铸质量异常中规则边界不清、工艺变量耦合强、根因标注不完整的问题，本文构建面向边界区失稳传播的三层监督框架。第一层通过高置信正负样本学习风险预警，第二层通过多标签异常集合描述复合质量现象，第三层利用 3xx/4xx 改钢原因与未改钢对照样本验证图路径解释的工艺一致性。该设计把预测性能、异常覆盖和解释可信度解耦，使模型优化不再依赖单一模糊标签。",
    )


def validate_docx(path: Path) -> None:
    doc = Document(path)
    table_text = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    text = "\n".join([p.text for p in doc.paragraphs] + table_text)
    required = [
        "当前任务重构与三数据集方案",
        "Risk-Binary Dataset",
        "Defect Multi-label Dataset",
        "Root-cause Alignment Dataset",
        "3xx、4xx 改钢原因样本",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"文档结构检查失败，缺少: {missing}")
    print(f"output={path}")
    print(f"size={path.stat().st_size}")
    print(f"paragraphs={len(doc.paragraphs)}")
    print(f"tables={len(doc.tables)}")


def main() -> None:
    src = choose_source()
    backup_existing_output()
    doc = Document(src)
    setup_styles(doc)
    append_three_dataset_plan(doc)
    doc.save(OUT)
    validate_docx(OUT)


if __name__ == "__main__":
    main()
