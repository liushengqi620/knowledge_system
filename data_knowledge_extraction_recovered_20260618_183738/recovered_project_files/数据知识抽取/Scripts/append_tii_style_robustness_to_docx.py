from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "数据知识挖掘.docx"
BACKUP_DIR = ROOT / "docs" / "word_backups"
ROBUST_DIR = ROOT / "knowledge_exports" / "major_quality_abnormality_robustness_v1"
ROBUST_JSON = ROBUST_DIR / "robustness_experiments_seed42.json"
ROBUST_SUMMARY = ROBUST_DIR / "robustness_summary.csv"
ROBUST_FIG_DIR = ROBUST_DIR / "figures"
DOC_FIG_DIR = ROBUST_DIR / "tii_docx_figures"

BLUE = RGBColor(31, 78, 121)
DARK = RGBColor(31, 77, 120)

PALETTE = {
    "blue": (54, 111, 196),
    "teal": (32, 145, 140),
    "orange": (245, 124, 0),
    "purple": (126, 87, 194),
    "green": (67, 160, 71),
    "red": (211, 47, 47),
    "light": (238, 242, 247),
    "grid": (218, 225, 235),
    "text": (32, 41, 55),
    "muted": (94, 107, 125),
    "bg": (255, 255, 255),
}


def font(size: int, *, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc") if bold else Path("C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/arial.ttf"),
    ]
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def setup_styles(doc: Document) -> None:
    for style_name in ("Normal", "List Bullet", "Caption"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run(run, size=10.5)


def add_bullet(doc: Document, text: str) -> None:
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
        prefix = ""
    except KeyError:
        paragraph = doc.add_paragraph()
        prefix = "- "
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(prefix + text)
    set_run(run, size=10.5)


def add_picture(doc: Document, path: Path, caption: str, *, width: float = 6.55) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_p = doc.add_paragraph()
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_p.paragraph_format.space_after = Pt(8)
    caption_run = caption_p.add_run(caption)
    set_run(caption_run, size=9.2, color=RGBColor(92, 100, 112))


def draw_text_fit(
    draw: ImageDraw.ImageDraw,
    xy: tuple[int, int],
    text: str,
    font_obj: ImageFont.ImageFont,
    *,
    max_width: int,
    fill: tuple[int, int, int] = PALETTE["text"],
) -> None:
    if draw.textlength(text, font=font_obj) <= max_width:
        draw.text(xy, text, font=font_obj, fill=fill)
        return
    clipped = text
    while clipped and draw.textlength(clipped + "...", font=font_obj) > max_width:
        clipped = clipped[:-1]
    draw.text(xy, clipped + "...", font=font_obj, fill=fill)


def rounded_box(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    body: str,
    *,
    fill: tuple[int, int, int],
    outline: tuple[int, int, int] = (210, 218, 230),
) -> None:
    x0, y0, x1, y1 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=outline, width=2)
    draw_text_fit(draw, (x0 + 24, y0 + 20), title, font(25, bold=True), max_width=x1 - x0 - 48)
    lines = body.split("\n")
    for idx, line in enumerate(lines):
        draw_text_fit(draw, (x0 + 24, y0 + 62 + idx * 30), line, font(19), max_width=x1 - x0 - 48, fill=PALETTE["muted"])


def arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], *, color: tuple[int, int, int] = PALETTE["muted"]) -> None:
    draw.line((start, end), fill=color, width=4)
    ex, ey = end
    sx, sy = start
    if ex >= sx:
        points = [(ex, ey), (ex - 16, ey - 9), (ex - 16, ey + 9)]
    else:
        points = [(ex, ey), (ex + 16, ey - 9), (ex + 16, ey + 9)]
    draw.polygon(points, fill=color)


def make_framework_figure(path: Path) -> None:
    img = Image.new("RGB", (1800, 1080), PALETTE["bg"])
    draw = ImageDraw.Draw(img)
    draw.text((70, 46), "KIEP-GL：边界区失稳传播驱动的品质异常识别框架", font=font(42, bold=True), fill=PALETTE["text"])
    draw.text((70, 104), "面向 IEEE TII 风格的方法图：数据重构、规则边界、事件时序、多头预测与路径解释统一建模", font=font(24), fill=PALETTE["muted"])
    draw.line((70, 150, 1730, 150), fill=PALETTE["grid"], width=2)

    boxes = [
        ((80, 220, 390, 390), "数据重构层", "品质异常代码去噪\n主要异常五类标签\n干净负样本与上下文分离", (238, 245, 255)),
        ((485, 220, 795, 390), "规则边界层", "rule_margin 距离\n上下限/量纲归一\n边界区敏感性分析", (238, 250, 248)),
        ((890, 220, 1200, 390), "事件时序层", "事件级 lag/delta\nrolling 统计\n失稳传播趋势", (255, 247, 236)),
        ((1295, 220, 1605, 390), "多头预测层", "主要异常二分类\n五类异常多标签\n阈值校准与消融验证", (244, 240, 255)),
    ]
    for box, title, body, fill in boxes:
        rounded_box(draw, box, title, body, fill=fill)
    arrow(draw, (390, 305), (485, 305))
    arrow(draw, (795, 305), (890, 305))
    arrow(draw, (1200, 305), (1295, 305))

    rounded_box(
        draw,
        (485, 525, 1200, 720),
        "工艺先验路径图",
        "变量节点 -> 工艺状态节点 -> 主要异常类别\n液面波动 -> 液面稳定性 -> 液面/卷渣\n南北拔热量差 -> 传热状态 -> 传热不均\n拉速/塞棒/流量 -> 流动控制 -> 控制异常",
        fill=(246, 248, 252),
    )
    arrow(draw, (1045, 390), (1045, 525))
    arrow(draw, (1295, 615), (1200, 615))

    rounded_box(
        draw,
        (80, 805, 520, 990),
        "实验验证",
        "Group split\n时间外推\n去重复派生特征消融\n多标签类别级评估",
        fill=(246, 248, 252),
    )
    rounded_box(
        draw,
        (680, 805, 1120, 990),
        "解释性验证",
        "Path Occlusion Drop\nRule Consistency@k\nPath Stability\nExpert Alignment",
        fill=(246, 248, 252),
    )
    rounded_box(
        draw,
        (1280, 805, 1720, 990),
        "论文贡献",
        "标签重构 + 边界失稳主线\n预测与溯源统一框架\n面向工业过程的鲁棒验证协议",
        fill=(246, 248, 252),
    )
    arrow(draw, (300, 720), (300, 805))
    arrow(draw, (840, 720), (900, 805))
    arrow(draw, (1450, 390), (1500, 805))
    img.save(path, quality=95)


def backup_docx() -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"数据知识挖掘.before_tii_style_robustness_{date.today().strftime('%Y%m%d')}.docx"
    if DOCX.exists() and not backup.exists():
        shutil.copy2(DOCX, backup)
    return backup


def metric(df: pd.DataFrame, task: str, experiment: str, name: str) -> float:
    value = df[(df["task"] == task) & (df["experiment"] == experiment)][name].iloc[0]
    return float(value)


def append_section(doc: Document, payload: dict[str, Any], summary: pd.DataFrame, figures: dict[str, Path]) -> None:
    today = date.today().isoformat()
    doc.add_page_break()
    add_heading(doc, f"面向 TII 投稿的方法框架与鲁棒性实验（{today}）", 1)
    add_para(
        doc,
        "本节按照 IEEE Transactions on Industrial Informatics（TII）论文常见写法，将当前工作整理为问题定义、方法框架、实验协议、鲁棒性结果和投稿注意事项。当前实验重点验证“主要异常标签重构 + 非泄漏工艺特征建模”的可靠性，并为后续加入 rule_margin、事件级时序传播和工艺路径解释提供可量化基线。",
    )

    add_heading(doc, "1. Problem Definition and Research Gap", 2)
    add_para(
        doc,
        "连铸品质异常预警的关键困难不在于单纯分类器能力不足，而在于标签边界模糊、异常代码多标签共现、工况状态与真实质量风险耦合、以及根因解释缺少可验证监督信号。若直接把所有品质异常代码混合作为单一类别，模型容易学习到不稳定的处置记录模式，而不是边界区工艺失稳传播机制。",
    )
    add_bullet(doc, "数据层面：将低支持代码和交接/质量状态类从主标签中拆出，构造主要异常、干净负样本和上下文辅助样本。")
    add_bullet(doc, "模型层面：采用“风险识别 + 异常多标签 + 工艺路径解释”的解耦式多头框架，避免所有目标被压进一个混杂分类头。")
    add_bullet(doc, "实验层面：除常规 group split 外，增加时间外推和去重复派生特征消融，检验模型是否依赖随机切分或重复字段。")

    add_heading(doc, "2. Proposed Framework", 2)
    make_framework_figure(figures["framework"])
    add_picture(doc, figures["framework"], "Fig. 1. Proposed KIEP-GL framework for boundary-region instability propagation modeling.", width=6.65)
    add_para(
        doc,
        "该框架的核心主线是“边界区失稳传播风险建模”。数据重构层提供清晰可靠的主要异常标签；规则边界层通过 rule_margin 描述工艺变量接近上下限的程度；事件时序层刻画边界扰动在连续窗口中的传播；多头预测层分别输出主要异常风险和五类异常现象；工艺先验路径图用于生成可解释溯源路径，并通过遮蔽实验和专家一致性进行验证。",
    )

    add_heading(doc, "3. Experimental Protocol", 2)
    add_bullet(doc, "Group split：按 process_sequence_key 分组切分，避免同一工况片段同时出现在训练集和测试集。")
    add_bullet(doc, "Time holdout：以 2022 年 1-8 月训练、9-10 月验证、11-12 月测试，模拟工业部署中的未来数据外推。")
    add_bullet(doc, "No-derived-alias ablation：移除 superheat、mold_level_range、cast_speed_range 等英文派生别名，仅保留原始工艺变量，检验指标是否由重复字段抬高。")
    add_bullet(doc, "Feature leakage control：排除品质异常代码、标签字段、改钢原因、处置代码、类别字段和样本 ID 字段，仅使用工艺数值变量。")
    add_para(
        doc,
        f"本轮鲁棒性实验使用 ExtraTreesClassifier，n_estimators={payload['n_estimators']}，随机种子为 {payload['seed']}，缺失值采用中位数填补。该设置不是最终复杂模型，而是论文中用于证明数据重构有效性和建立非泄漏强基线的基础模型。",
    )

    add_heading(doc, "4. Robustness Results", 2)
    b_group = metric(summary, "binary", "group_all_features", "f1")
    b_time = metric(summary, "binary", "time_all_features", "f1")
    b_noalias = metric(summary, "binary", "time_no_derived_alias", "f1")
    m_group = metric(summary, "multilabel", "group_all_features", "macro_f1")
    m_time = metric(summary, "multilabel", "time_all_features", "macro_f1")
    m_noalias = metric(summary, "multilabel", "time_no_derived_alias", "macro_f1")
    add_para(
        doc,
        f"结果显示，主要异常二分类在 group split 下 F1={b_group:.3f}，时间外推下 F1={b_time:.3f}；进一步移除英文派生别名后，最严格设置下 F1 仍达到 {b_noalias:.3f}。五类主要异常多标签识别在 group split 下 Macro-F1={m_group:.3f}，时间外推下 Macro-F1={m_time:.3f}，移除派生别名后时间外推 Macro-F1={m_noalias:.3f}。",
    )
    add_picture(doc, figures["binary_metrics"], "Fig. 2. Robustness results for binary major-abnormality warning.", width=6.55)
    add_picture(doc, figures["multilabel_metrics"], "Fig. 3. Robustness results for five-class multilabel abnormality recognition.", width=6.55)
    add_picture(doc, figures["generalization_drop"], "Fig. 4. Generalization gap and feature de-duplication ablation.", width=6.35)

    add_heading(doc, "5. Innovation Claims for a TII-Style Manuscript", 2)
    add_bullet(doc, "Label-space innovation：从原始品质异常代码中构造主要异常多标签任务，并将交接/质量状态类作为上下文而非主标签，降低标签噪声。")
    add_bullet(doc, "Problem-formulation innovation：将任务定义为边界区失稳传播风险建模，而不是普通缺陷分类，使 rule_margin、时序传播和路径解释服务于同一主线。")
    add_bullet(doc, "Modeling innovation：采用风险识别、多标签异常现象识别和工艺先验路径解释的解耦式架构，避免复杂模块简单拼接。")
    add_bullet(doc, "Validation innovation：引入时间外推和去重复派生特征消融，证明模型不是依赖随机切分相似性或重复字段获得高分。")
    add_bullet(doc, "Interpretability roadmap：后续应补充 Path Occlusion Drop、Rule Consistency@k、Path Stability 和 Expert Alignment，形成预测性能与机理解释并重的证据链。")

    add_heading(doc, "6. Submission Notes and Remaining Evidence", 2)
    add_para(
        doc,
        "若以 TII 或同级工业信息学期刊为目标，当前结果已经适合作为数据重构和强基线证据，但还不能直接作为完整投稿版本。下一轮必须补足 rule_margin 增益实验、事件级时序传播消融、工艺路径解释定量评价和多随机种子置信区间。特别是传热不均类样本较少，需要单独报告支持数、置信区间和类别稳定性，避免审稿人质疑少样本类别的指标偶然性。",
    )


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not ROBUST_JSON.exists():
        raise FileNotFoundError(ROBUST_JSON)
    backup = backup_docx()
    DOC_FIG_DIR.mkdir(parents=True, exist_ok=True)
    figures = {
        "framework": DOC_FIG_DIR / "tii_method_framework.png",
        "binary_metrics": ROBUST_FIG_DIR / "robustness_binary_metrics.png",
        "multilabel_metrics": ROBUST_FIG_DIR / "robustness_multilabel_metrics.png",
        "generalization_drop": ROBUST_FIG_DIR / "robustness_generalization_drop.png",
    }
    payload = json.loads(ROBUST_JSON.read_text(encoding="utf-8"))
    summary = pd.read_csv(ROBUST_SUMMARY)
    doc = Document(DOCX)
    setup_styles(doc)
    append_section(doc, payload, summary, figures)
    doc.save(DOCX)
    print(f"updated={DOCX}")
    print(f"backup={backup}")
    print(f"framework={figures['framework']}")


if __name__ == "__main__":
    main()
