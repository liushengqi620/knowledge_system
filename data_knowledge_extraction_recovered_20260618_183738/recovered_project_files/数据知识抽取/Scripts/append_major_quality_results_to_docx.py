from __future__ import annotations

import json
import shutil
from datetime import date
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "数据知识挖掘.docx"
BACKUP_DIR = ROOT / "docs" / "word_backups"
RESULT_JSON = (
    ROOT
    / "knowledge_exports"
    / "major_quality_abnormality_experiment_v1"
    / "major_quality_abnormality_experiment_seed42.json"
)
BINARY_IMPORTANCE = (
    ROOT
    / "knowledge_exports"
    / "major_quality_abnormality_experiment_v1"
    / "binary_feature_importance_seed42.csv"
)
MULTILABEL_IMPORTANCE = (
    ROOT
    / "knowledge_exports"
    / "major_quality_abnormality_experiment_v1"
    / "multilabel_feature_importance_seed42.csv"
)

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
CALLOUT_FILL = "F4F6F9"


def fmt(value: Any, digits: int = 3) -> str:
    if isinstance(value, float):
        return f"{value:.{digits}f}"
    return str(value)


def set_run(run, size: float | None = None, bold: bool | None = None, color: RGBColor | None = None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def setup_styles(doc: Document) -> None:
    for style_name in ("Normal", "List Bullet"):
        if style_name in [s.name for s in doc.styles]:
            style = doc.styles[style_name]
            style.font.name = "Calibri"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            style.font.size = Pt(10.5)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), bold=True, color=BLUE if level <= 2 else DARK)


def add_para(doc: Document, text: str, *, bold_label: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    if bold_label:
        run = paragraph.add_run(bold_label)
        set_run(run, bold=True)
        run = paragraph.add_run(text)
        set_run(run)
    else:
        run = paragraph.add_run(text)
        set_run(run)


def add_bullet(doc: Document, text: str) -> None:
    try:
        paragraph = doc.add_paragraph(style="List Bullet")
        prefix = ""
    except KeyError:
        paragraph = doc.add_paragraph()
        prefix = "- "
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(prefix + text)
    set_run(run, size=10.5)


def add_callout(doc: Document, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, CALLOUT_FILL)
    margins(cell, top=120, bottom=120, start=160, end=160)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    set_run(run, bold=True, color=DARK)
    run = paragraph.add_run("\n" + body)
    set_run(run)
    doc.add_paragraph()


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade(cell, GRAY_FILL)
        margins(cell)
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run(run, size=9.5, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            margins(cell)
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            set_run(run, size=9)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    doc.add_paragraph()


def backup_docx() -> Path:
    BACKUP_DIR.mkdir(parents=True, exist_ok=True)
    backup = BACKUP_DIR / f"数据知识挖掘.before_major_quality_results_{date.today().strftime('%Y%m%d')}.docx"
    if not backup.exists():
        shutil.copy2(DOCX, backup)
    return backup


def load_result() -> dict[str, Any]:
    return json.loads(RESULT_JSON.read_text(encoding="utf-8"))


def append_section(doc: Document, result: dict[str, Any]) -> None:
    binary = result["binary"]
    multilabel = result["multilabel"]
    binary_imp = pd.read_csv(BINARY_IMPORTANCE).head(10)
    multilabel_imp = pd.read_csv(MULTILABEL_IMPORTANCE).head(10)

    today = date.today().isoformat()
    doc.add_page_break()
    add_heading(doc, f"主要异常数据处理与模型基线结果（{today}）", 1)
    add_callout(
        doc,
        "本轮结论",
        "将品质异常代码重新收敛为“主要异常”后，模型效果明显提升。当前结果说明此前性能偏低的主要原因并非模型能力不足，而是标签口径过宽、交接状态类与低支持代码混入主任务，导致正负样本边界不稳定。新方案把交接过渡/质量状态类作为上下文辅助变量，而不是主标签，从而使风险识别和异常类型识别更加清晰。",
    )

    add_heading(doc, "1. 数据处理口径", 2)
    add_para(
        doc,
        "本轮数据入口为 knowledge_exports/major_quality_abnormality_v1。主要异常被限定为五类：液面/卷渣质量类、温度/保护剂质量类、传热不均质量类、拉速-塞棒/流量控制异常类、过程参数异常类。交接过渡/质量状态类不作为主训练标签，只作为工况上下文和后续消融分析变量保留。",
    )
    add_table(
        doc,
        ["数据项", "规模", "用途"],
        [
            ["主要异常样本", "45,185", "五类主要异常多标签识别"],
            ["干净负样本", "30,187", "主要异常二分类中的稳定负样本"],
            ["辅助状态-only 样本", "23,374", "作为交接/质量状态上下文，不进入主标签"],
            ["二分类训练集", "75,372", "主要异常 vs 干净负样本"],
            ["主异常多标签样本", "9,742（21.56%）", "说明异常共现明显，应优先采用多标签任务"],
        ],
        [1.55, 1.55, 3.20],
    )

    add_heading(doc, "2. 模型设置", 2)
    add_bullet(doc, "模型采用 ExtraTreesClassifier，80 棵树，随机种子 42，中位数填补缺失值。")
    add_bullet(doc, "切分方式采用 process_sequence_key 分组切分，降低同一工况片段同时出现在训练集和测试集的风险。")
    add_bullet(doc, "特征策略为仅使用 99 个工艺数值特征；显式排除品质异常代码、标签字段、改钢原因、处置代码、类别字段和样本 ID 字段，避免直接标签泄漏。")
    add_bullet(doc, "本轮结果是快速诊断基线，后续仍需要增加时间外推、连铸线外推和重复种子验证。")

    add_heading(doc, "3. 主要异常二分类结果", 2)
    add_table(
        doc,
        ["指标", "结果"],
        [
            ["F1", fmt(binary["f1"])],
            ["AUC", fmt(binary["roc_auc"])],
            ["AP", fmt(binary["average_precision"])],
            ["Accuracy", fmt(binary["accuracy"])],
            ["Precision", fmt(binary["precision"])],
            ["Recall", fmt(binary["recall"])],
            ["Balanced Accuracy", fmt(binary["balanced_accuracy"])],
            ["混淆矩阵", "TN=5,595，FP=310，FN=462，TP=8,708"],
        ],
        [2.10, 3.20],
    )

    add_heading(doc, "4. 五类主要异常多标签结果", 2)
    add_table(
        doc,
        ["总体指标", "结果"],
        [
            ["Micro-F1", fmt(multilabel["micro_f1"])],
            ["Macro-F1", fmt(multilabel["macro_f1"])],
            ["mAP", fmt(multilabel["mean_average_precision"])],
            ["Samples-F1", fmt(multilabel["sample_f1"])],
            ["Subset Accuracy", fmt(multilabel["subset_accuracy"])],
        ],
        [2.10, 3.20],
    )
    per_rows: list[list[str]] = []
    for label_key, item in multilabel["per_label"].items():
        per_rows.append(
            [
                item["display_name"],
                str(item["support"]),
                fmt(item["precision"]),
                fmt(item["recall"]),
                fmt(item["f1"]),
                fmt(item["average_precision"]),
            ]
        )
    add_table(
        doc,
        ["异常类别", "测试支持数", "Precision", "Recall", "F1", "AP"],
        per_rows,
        [1.75, 0.80, 0.75, 0.75, 0.75, 0.75],
    )

    add_heading(doc, "5. 关键工艺特征与机理一致性", 2)
    add_para(
        doc,
        "特征重要性前列集中在过热度、TD 温度、液面波动、拉速波动、TD 吨位波动、南北拔热量差等工艺变量，与连铸品质异常机理基本一致。尤其是液面/卷渣类由液面波动主导，传热不均类由南北拔热量差和 heat_exchange_ns_diff 主导，流量控制类由拉速波动和 TD_weight_range 主导，说明模型并非简单依赖异常编码字段。",
    )
    add_table(
        doc,
        ["二分类重要特征", "重要性"],
        [[str(row["feature"]), fmt(float(row["binary_importance"]), 4)] for _, row in binary_imp.iterrows()],
        [3.55, 1.00],
    )
    add_table(
        doc,
        ["多标签平均重要特征", "平均重要性"],
        [
            [str(row["feature"]), fmt(float(row["multilabel_mean_importance"]), 4)]
            for _, row in multilabel_imp.iterrows()
        ],
        [3.55, 1.00],
    )

    add_heading(doc, "6. 结果解释与下一步验证", 2)
    add_bullet(doc, "当前结果表明重新定义标签后，数据与模型适配度显著提升；此前效果差主要来自标签混杂和负样本不稳定，而不是单纯模型结构不足。")
    add_bullet(doc, "结果较高，需要进一步做更严格验证：按月份做时间外推，按连铸线/工况做外推，去掉中英文重复派生特征后重跑。")
    add_bullet(doc, "下一步模型优化应引入 rule_margin、非线性时序 lag 特征和工艺先验路径遮蔽验证，用来证明规则边界、时序传播和路径解释确实带来增益。")
    add_bullet(doc, "传热不均类测试支持数只有 74，虽然 F1 较高，但论文中应单独说明样本量限制，并报告置信区间或多 seed 稳定性。")


def main() -> None:
    if not DOCX.exists():
        raise FileNotFoundError(DOCX)
    if not RESULT_JSON.exists():
        raise FileNotFoundError(RESULT_JSON)
    backup = backup_docx()
    doc = Document(DOCX)
    setup_styles(doc)
    append_section(doc, load_result())
    doc.save(DOCX)
    print(f"updated={DOCX}")
    print(f"backup={backup}")


if __name__ == "__main__":
    main()
